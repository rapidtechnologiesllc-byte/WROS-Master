"""
Salary Structure Generator — Updated Formulas
==============================================
CTC breakdown rules:
  Basic           = 50% of CTC
  HRA             = 40% of Basic
  Medical         = ₹15,000 fixed
  Transport       = ₹19,200 fixed
  Deployment      = min(CTC - Basic - HRA - Medical - Transport, 60,000)
  Fixed Allowance = CTC - Basic - HRA - Medical - Transport - Deployment

EPF base (excl. HRA + Deployment) = Basic + Transport + Medical + Fixed
  EPF Employee = min(EPF_base * 12%, 21,600)
  EPF Employer = min(EPF_base * 12%, 21,600)

ESIC base = same as EPF base  (only if ESIC_gross_pm ≤ ₹21,000)
  ESIC Employee = 0.75% of ESIC_base
  ESIC Employer = 3.25% of ESIC_base  ← informational / other-benefit

Total Deductions = EPF_employee + EPF_employer + ESIC_employee
Net Income       = CTC - Total Deductions
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Salary maths
# ---------------------------------------------------------------------------

@dataclass
class SalaryBreakdown:
    employee_name: str
    annual_ctc: float
 
    # Earnings (annual)
    basic_pa: float = 0.0
    hra_pa: float = 0.0
    medical_pa: float = 15_000.0
    transport_pa: float = 19_200.0
    deployment_pa: float = 0.0        # capped at 60,000; NOT included in EPF/ESIC base
    fixed_allowance_pa: float = 0.0   # remainder after all other components
 
    # Deductions (annual)
    epf_employee_pa: float = 0.0
    epf_employer_pa: float = 0.0
    esic_employee_pa: float = 0.0
    esic_employer_pa: float = 0.0     # informational — shown as "Other Benefit", not deducted
 
    # ── derived ──────────────────────────────────────────────────────────────
    @property
    def gross_pa(self) -> float:
        """Total CTC = sum of all earning components."""
        return (
            self.basic_pa
            + self.hra_pa
            + self.medical_pa
            + self.transport_pa
            + self.deployment_pa
            + self.fixed_allowance_pa
        )
 
    @property
    def total_deductions_pa(self) -> float:
        """
        Excel 'Total Deductions (B)':
          EPF Employee + EPF Employer + ESIC Employee.
        ESIC Employer is an *other benefit* (not deducted from gross).
        """
        return self.epf_employee_pa + self.epf_employer_pa + self.esic_employee_pa
 
    @property
    def net_pa(self) -> float:
        """Net Income (A - B = C) per Excel."""
        return self.gross_pa - self.total_deductions_pa
 
    # ── monthly helpers ───────────────────────────────────────────────────────
    @property
    def basic_pm(self):             return self.basic_pa / 12
    @property
    def hra_pm(self):               return self.hra_pa / 12
    @property
    def medical_pm(self):           return self.medical_pa / 12
    @property
    def transport_pm(self):         return self.transport_pa / 12
    @property
    def deployment_pm(self):        return self.deployment_pa / 12
    @property
    def fixed_allowance_pm(self):   return self.fixed_allowance_pa / 12
    @property
    def gross_pm(self):             return self.gross_pa / 12
    @property
    def epf_employee_pm(self):      return self.epf_employee_pa / 12
    @property
    def epf_employer_pm(self):      return self.epf_employer_pa / 12
    @property
    def esic_employee_pm(self):     return self.esic_employee_pa / 12
    @property
    def esic_employer_pm(self):     return self.esic_employer_pa / 12
    @property
    def total_deductions_pm(self):  return self.total_deductions_pa / 12
    @property
    def net_pm(self):               return self.net_pa / 12
 
 
def calculate_salary(employee_name: str, annual_ctc: float) -> SalaryBreakdown:
    """
    Compute all salary components from annual CTC per the Excel salary structure.
 
    Component rules (all figures annual unless noted):
    ┌─────────────────────┬────────────────────────────────────────────────┐
    │ Basic               │ 50% of CTC                                     │
    │ HRA                 │ 40% of Basic                                   │
    │ Medical Allowance   │ Fixed ₹15,000                                  │
    │ Transport Allowance │ Fixed ₹19,200                                  │
    │ Deployment Allow.   │ min(remaining after above 4, ₹60,000); ≥ 0    │
    │ Fixed Allowance     │ CTC − (Basic+HRA+Medical+Transport+Deployment) │
    └─────────────────────┴────────────────────────────────────────────────┘
 
    EPF/ESIC base = Basic + Transport + Medical + Fixed Allowance
                    (HRA and Deployment Allowance excluded per Excel spec)
 
    EPF (Employee & Employer each):
        IF epf_esic_base > ₹1,80,000 (annual)  →  ₹21,600 (cap)
        ELSE                                    →  12% × epf_esic_base
 
    ESIC (applies only when monthly epf_esic_base ≤ ₹21,000):
        Employee share  = 0.75% × epf_esic_base
        Employer share  = 3.25% × epf_esic_base  (other benefit, not deducted)
    """
    # ── Earnings ──────────────────────────────────────────────────────────────
    basic     = annual_ctc * 0.50
    hra       = basic * 0.40
    medical   = 15_000.0
    transport = 19_200.0
 
    # Deployment: what remains after the first four components, capped at ₹60,000
    remaining_after_four = annual_ctc - basic - hra - medical - transport
    deployment = min(remaining_after_four, 60_000.0)
    deployment = max(deployment, 0.0)   # guard for very low CTCs
 
    # Fixed allowance: whatever is left over after deployment
    fixed = annual_ctc - basic - hra - medical - transport - deployment
    fixed = max(fixed, 0.0)             # guard for very low CTCs
 
    # ── EPF Employee Base ─────────────────────────────────────────────────────
    # Includes: Basic + Transport + Medical + Fixed
    # Excludes: HRA, Deployment
    epf_employee_base = basic + transport + medical + fixed

    # ── EPF Employer Base ─────────────────────────────────────────────────────
    # Includes: Basic + Transport + Medical + Fixed
    # Excludes: HRA, Deployment
    epf_employer_base = basic + transport + medical + fixed

    # ── EPF Constants ─────────────────────────────────────────────────────────
    EPF_ANNUAL_CEILING = 180_000.0   # ₹15,000/month × 12
    EPF_CAP            = 21_600.0    # 12% of ₹1,80,000

    # ── EPF Employee ──────────────────────────────────────────────────────────
    if epf_employee_base > EPF_ANNUAL_CEILING:
        epf_employee = EPF_CAP
    else:
        epf_employee = epf_employee_base * 0.12

    # ── EPF Employer ──────────────────────────────────────────────────────────
    if epf_employer_base > EPF_ANNUAL_CEILING:
        epf_employer = EPF_CAP
    else:
        epf_employer = epf_employer_base * 0.12

    # ── ESIC Base (same as EPF Employee Base) ─────────────────────────────────
    esic_base = epf_employee_base

    # ── ESIC ──────────────────────────────────────────────────────────────────
    ESIC_MONTHLY_THRESHOLD = 21_000.0
    if (esic_base / 12) <= ESIC_MONTHLY_THRESHOLD:
        esic_employee = esic_base * 0.0075   # 0.75% — deducted from employee
        esic_employer = esic_base * 0.0325   # 3.25% — shown as Other Benefit
    else:
        esic_employee = 0.0
        esic_employer = 0.0

    # ── Round to 2 decimal places ─────────────────────────────────────────────
    epf_employee  = round(epf_employee,  2)
    epf_employer  = round(epf_employer,  2)
    esic_employee = round(esic_employee, 2)
    esic_employer = round(esic_employer, 2)

    return SalaryBreakdown(
        employee_name=employee_name,
        annual_ctc=annual_ctc,
        basic_pa=basic,
        hra_pa=hra,
        medical_pa=medical,
        transport_pa=transport,
        deployment_pa=deployment,
        fixed_allowance_pa=fixed,
        epf_employee_pa=epf_employee,
        epf_employer_pa=epf_employer,
        esic_employee_pa=esic_employee,
        esic_employer_pa=esic_employer,
    )


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

_DARK_BLUE  = "1F4E79"
_MID_BLUE   = "2E75B6"
_LIGHT_BLUE = "D6E4F0"
_WHITE      = "FFFFFF"
_GOLD       = "C9A84C"
_LIGHT_GREY = "F2F2F2"


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _write_cell(
    cell,
    text: str,
    bold: bool = False,
    font_size: int = 9,
    color: Optional[str] = None,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def generate_salary_structure_docx(
    employee_name: str,
    annual_ctc: float,
) -> bytes:
    """Generate and return a salary-structure Word document as raw bytes."""
    sal = calculate_salary(employee_name, annual_ctc)
    doc = Document()

    # page margins
    for sec in doc.sections:
        from docx.shared import Cm
        sec.top_margin = sec.bottom_margin = Cm(1.5)
        sec.left_margin = sec.right_margin = Cm(2.0)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BlitzenX Solutions Private Limited")
    r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(_DARK_BLUE)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("COMPENSATION & BENEFITS STATEMENT")
    r2.bold = True; r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor.from_string(_MID_BLUE)

    doc.add_paragraph()

    info = doc.add_paragraph()
    b = info.add_run("Employee Name: "); b.bold = True
    info.add_run(sal.employee_name)
    info.add_run("     |     ")
    b2 = info.add_run("Annual CTC: "); b2.bold = True
    info.add_run(f"₹ {sal.annual_ctc:,.2f}")
    info.paragraph_format.space_after = Pt(10)
    doc.add_paragraph()

    R = WD_ALIGN_PARAGRAPH.RIGHT
    C = WD_ALIGN_PARAGRAPH.CENTER
    L = WD_ALIGN_PARAGRAPH.LEFT

    # --- build table ---
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    col_w = [Inches(3.0), Inches(2.0), Inches(1.3), Inches(1.3)]

    def add_row(texts, bg=_WHITE, bolds=None, fg=None, aligns=None, fs=9):
        row    = table.add_row()
        bolds  = bolds  or [False] * 4
        fg     = fg     or [None]  * 4
        aligns = aligns or [L]     * 4
        for i, (cell, txt) in enumerate(zip(row.cells, texts)):
            _set_cell_bg(cell, bg)
            _write_cell(cell, txt, bold=bolds[i], font_size=fs,
                        color=fg[i], align=aligns[i])
            cell.width = col_w[i]

    W = [_WHITE] * 4

    # Header
    add_row(
        ["Components of the Salary Package", "Remarks", "Amount (P.M.)", "Amount (P.A.)"],
        bg=_DARK_BLUE, bolds=[True]*4, fg=[_WHITE]*4,
        aligns=[L, C, R, R],
    )

    # Earnings section title
    add_row(["Basic Salary Components", "", "", ""],
            bg=_MID_BLUE, bolds=[True,False,False,False], fg=[_WHITE]+[None]*3)

    earnings = [
        ("Basic Salary",                          "50% of CTC",                          sal.basic_pm,           sal.basic_pa),
        ("House Rent Allowance (HRA)",             "40% of Basic Salary",                 sal.hra_pm,             sal.hra_pa),
        ("Medical Allowances",                     "Fixed ₹15,000 p.a.",                  sal.medical_pm,         sal.medical_pa),
        ("Transport Allowances",                   "Fixed ₹19,200 p.a.",                  sal.transport_pm,       sal.transport_pa),
        ("Deployment / Performance Allowances",    "Fixed, capped ₹60,000 p.a.",          sal.deployment_pm,      sal.deployment_pa),
        ("Fixed Allowances",                       "Remaining amount after above",         sal.fixed_allowance_pm, sal.fixed_allowance_pa),
    ]
    for i, (name, note, pm, pa) in enumerate(earnings):
        bg = _LIGHT_BLUE if i % 2 == 0 else _WHITE
        add_row([name, note, f"{pm:,.2f}", f"{pa:,.2f}"], bg=bg, aligns=[L, L, R, R])

    add_row(
        ["Gross / CTC  (A)", "", f"{sal.gross_pm:,.2f}", f"{sal.gross_pa:,.2f}"],
        bg=_MID_BLUE, bolds=[True,False,True,True],
        fg=[_WHITE,None,_WHITE,_WHITE], aligns=[L,C,R,R], fs=10,
    )

    # Spacer
    add_row(["", "", "", ""])

    # Deductions section
    add_row(["Less Deductions:", "", "", ""],
            bg=_DARK_BLUE, bolds=[True,False,False,False], fg=[_WHITE]+[None]*3)

    deductions = [
        ("EPF — Employee Contribution (12%)",  "Excl. HRA & Deployment; capped ₹21,600", sal.epf_employee_pm, sal.epf_employee_pa),
        ("EPF — Employer Contribution (12%)",  "Excl. HRA & Deployment; capped ₹21,600", sal.epf_employer_pm, sal.epf_employer_pa),
        ("ESIC — Employee (0.75%)",            "If ESIC gross ≤ ₹21,000/month",          sal.esic_employee_pm, sal.esic_employee_pa),
    ]
    for i, (name, note, pm, pa) in enumerate(deductions):
        bg = _LIGHT_BLUE if i % 2 == 0 else _WHITE
        add_row([name, note, f"{pm:,.2f}", f"{pa:,.2f}"], bg=bg, aligns=[L, L, R, R])

    add_row(
        ["Total Deductions  (B)", "", f"{sal.total_deductions_pm:,.2f}", f"{sal.total_deductions_pa:,.2f}"],
        bg=_MID_BLUE, bolds=[True,False,True,True],
        fg=[_WHITE,None,_WHITE,_WHITE], aligns=[L,C,R,R],
    )

    add_row(["", "", "", ""])

    # Net income
    add_row(
        ["Net Income  (C = A − B)", "", f"{sal.net_pm:,.2f}", f"{sal.net_pa:,.2f}"],
        bg=_GOLD, bolds=[True,False,True,True],
        fg=[_WHITE,None,_WHITE,_WHITE], aligns=[L,C,R,R], fs=11,
    )

    # Other Benefits
    doc.add_paragraph()
    bt = doc.add_paragraph()
    br = bt.add_run("Other Benefits Costed  (select as applicable per employee)")
    br.bold = True; br.font.size = Pt(11)
    br.font.color.rgb = RGBColor.from_string(_DARK_BLUE)

    b_table = doc.add_table(rows=0, cols=3)
    b_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    b_table.style = "Table Grid"
    b_col_w = [Inches(3.0), Inches(1.6), Inches(1.6)]

    def add_b_row(texts, bg=_WHITE, bolds=None, fg=None):
        row    = b_table.add_row()
        bolds  = bolds or [False]*3
        fg     = fg    or [None]*3
        for i, (cell, txt) in enumerate(zip(row.cells, texts)):
            _set_cell_bg(cell, bg)
            _write_cell(cell, txt, bold=bolds[i], font_size=9,
                        color=fg[i],
                        align=R if i > 0 else L)
            cell.width = b_col_w[i]

    add_b_row(["Benefit", "Monthly (₹)", "Annual (₹)"],
              bg=_DARK_BLUE, bolds=[True]*3, fg=[_WHITE]*3)

    esic_emp_pm = sal.esic_employer_pm
    esic_emp_pa = sal.esic_employer_pa

    other = [
        ("ESIC — Employer (3.25%)",                          esic_emp_pm,    esic_emp_pa),
        ("Health Insurance (Employee, Spouse & Children)",   0,      25_600.0),
        ("Guidewire Certification",                          0,     212_500.0),
        ("Paid Time Off",                                    0,     sal.gross_pa/12),
        ("Accidental Policy",                                714.0,          1_440.0),
        ("Term Insurance",                                  1_650.0,        3_000.0),
    ]
    for i, (name, pm, pa) in enumerate(other):
        bg = _LIGHT_BLUE if i % 2 == 0 else _WHITE
        add_b_row([name, f"{pm:,.2f}", f"{pa:,.2f}"], bg=bg)

    total_ob = sum(pa for _, _, pa in other)
    add_b_row(
        ["Total Other Benefits", "—", f"{total_ob:,.2f}"],
        bg=_MID_BLUE, bolds=[True]*3, fg=[_WHITE]*3,
    )

    # Footer
    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run(
        "Note: System-generated document. All figures in Indian Rupees (₹). "
        "EPF & ESIC calculated per statutory rules. "
        "Other benefits are informational and selected per employee requirements."
    )
    nr.italic = True; nr.font.size = Pt(8)
    nr.font.color.rgb = RGBColor.from_string("777777")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def get_salary_filename(employee_name: str) -> str:
    safe = employee_name.strip().replace(" ", "_").replace("/", "-")
    return f"Salary_Structure_{safe}.docx"
