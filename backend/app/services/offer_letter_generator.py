"""
Offer Letter Generator
======================
Fetches the .docx template from SharePoint, replaces {{placeholder}} tokens
import logging
with real candidate / offer data, and returns the filled document as bytes.

Placeholder convention (matching the BlitzenX template exactly):
  {{BasicInfo.FirstName}}
  {{BasicInfo.LastName}}
  {{BasicInfo.JobTitle}}
  {{BasicInfo.FullName}}
  {{BasicInfo.Department}}
  {{BasicInfo.Location}}
  {{BasicInfo.DateOfferExpires}}
  {{BasicInfo.MediumDate}}               — letter date (today)
  {{BasicInfo.DateJoiningMedium}}        — joining date
  {{BasicInfo.AnnualSalary}}
  {{BasicInfo.FullSalaryWithBonusOthersWithoutDeduction}}
  {{BasicInfo.FullSalaryStructureWithCTCDeductionsAndNetPay}}
                                         ↑ replaced with a full salary table
  {{JobInfo.Location}}
  {{Signature SignatureHRManager2}}      — hiring manager name
  {{Signature CandidateSignature}}       — candidate name (signature line)
"""

import io
import os
import re
from dataclasses import dataclass
from datetime import date, datetime
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.core.logging import logger
from app.services.sharepoint_service import download_file
from app.services.salary_structure_generator import calculate_salary


# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
TEMPLATE_PATH = os.getenv(
    "SHAREPOINT_TEMPLATE_PATH",
    "templates/Internship Offer letter.docx",
)
FULLTIME_TEMPLATE_PATH = os.getenv(
    "SHAREPOINT_FULLTIME_TEMPLATE_PATH",
    "templates/Full Time Offer letter.docx",
)
GENERATED_FOLDER = os.getenv(
    "SHAREPOINT_GENERATED_FOLDER",
    "generated-offers",
)
HR_MANAGER_SIGNATURE_PATH = os.getenv(
    "SHAREPOINT_HR_MANAGER_SIGNATURE_PATH",
    "templates/hiringmanager_signature.png",
)

# Friendly name → SharePoint template path mapping
_TEMPLATE_MAP: dict[str, str] = {
    "intern":    TEMPLATE_PATH,
    "fulltime":  FULLTIME_TEMPLATE_PATH,
}


def get_template_path(template_type: str) -> str:
    """
    Resolve a user-friendly template type string to the corresponding
    SharePoint path.

    Args:
        template_type: "intern" or "fulltime" (case-insensitive).

    Returns:
        SharePoint-relative path string.

    Raises:
        ValueError: If the type is not recognised.
    """
    key = template_type.strip().lower()
    if key not in _TEMPLATE_MAP:
        valid = ", ".join(f"'{k}'" for k in _TEMPLATE_MAP)
        raise ValueError(
            f"Unknown template_type '{template_type}'. Valid options: {valid}."
        )
    return _TEMPLATE_MAP[key]

# Sentinel used internally — the salary placeholder is replaced with this
# string first, then found again and swapped out for a real table.
_SALARY_TABLE_MARKER = "__BLITZENX_SALARY_TABLE__"

# Colour palette (matches salary_structure_generator.py)
_C_DARK_BLUE  = "1F4E79"
_C_MID_BLUE   = "2E75B6"
_C_LIGHT_BLUE = "D6E4F0"
_C_WHITE      = "FFFFFF"
_C_GOLD       = "C9A84C"


# ---------------------------------------------------------------------------
# Date helpers
# ---------------------------------------------------------------------------

_MONTH_NAMES = [
    "", "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]


def _fmt_date_medium(d) -> str:
    """Format a date/datetime as '30 April 2026'."""
    if d is None:
        return "TBD"
    if isinstance(d, datetime):
        d = d.date()
    return f"{d.day} {_MONTH_NAMES[d.month]} {d.year}"


# ---------------------------------------------------------------------------
# Salary maths
# ---------------------------------------------------------------------------

@dataclass
class _Salary:
    annual_ctc:         float
    basic_pa:           float
    hra_pa:             float
    medical_pa:         float = 15_000.0
    transport_pa:       float = 19_200.0
    deployment_pa:      float = 0.0
    fixed_allowance_pa: float = 0.0
    epf_employee_pa:    float = 0.0
    epf_employer_pa:    float = 0.0
    esic_employee_pa:   float = 0.0

    @property
    def gross_pa(self):
        return (self.basic_pa + self.hra_pa + self.medical_pa
                + self.transport_pa + self.deployment_pa
                + self.fixed_allowance_pa)
    @property
    def total_deductions_pa(self):
        return self.epf_employee_pa + self.epf_employer_pa + self.esic_employee_pa
    @property
    def net_pa(self):         return self.gross_pa - self.total_deductions_pa
    @property
    def basic_pm(self):       return self.basic_pa / 12
    @property
    def hra_pm(self):         return self.hra_pa / 12
    @property
    def medical_pm(self):     return self.medical_pa / 12
    @property
    def transport_pm(self):   return self.transport_pa / 12
    @property
    def deployment_pm(self):  return self.deployment_pa / 12
    @property
    def fixed_pm(self):       return self.fixed_allowance_pa / 12
    @property
    def gross_pm(self):       return self.gross_pa / 12
    @property
    def epf_emp_pm(self):     return self.epf_employee_pa / 12
    @property
    def epf_er_pm(self):      return self.epf_employer_pa / 12
    @property
    def esic_emp_pm(self):    return self.esic_employee_pa / 12
    @property
    def total_ded_pm(self):   return self.total_deductions_pa / 12
    @property
    def net_pm(self):         return self.net_pa / 12


def _calc_salary(annual_ctc: float) -> _Salary:
    basic     = annual_ctc * 0.50
    hra       = basic * 0.40
    medical   = 15_000.0
    transport = 19_200.0
    deployment = min(max(annual_ctc - basic - hra - medical - transport, 0.0), 60_000.0)
    fixed      = max(annual_ctc - basic - hra - medical - transport - deployment, 0.0)

    # EPF/ESIC base: Basic + Transport + Medical + Fixed  (NO HRA, NO Deployment)
    epf_base = basic + transport + medical + fixed

    # EPF: if base > ₹1,80,000 → cap at ₹21,600; else 12% of base
    if epf_base > 1_80_000.0:
        epf_emp = 21_600.0
        epf_er  = 21_600.0
    else:
        epf_emp = epf_base * 0.12
        epf_er  = epf_base * 0.12

    # ESIC: only when monthly EPF/ESIC base ≤ ₹21,000
    if (epf_base / 12) <= 21_000.0:
        esic_emp = epf_base * 0.0075   # employee 0.75% — deducted
    else:
        esic_emp = 0.0

    return _Salary(
        annual_ctc=annual_ctc, basic_pa=basic, hra_pa=hra,
        medical_pa=medical, transport_pa=transport,
        deployment_pa=deployment, fixed_allowance_pa=fixed,
        epf_employee_pa=epf_emp, epf_employer_pa=epf_er,
        esic_employee_pa=esic_emp,
    )


def _parse_salary(salary_str: str) -> float:
    """Parse a salary string like '3,60,000' or '₹360000' to float."""
    try:
        cleaned = str(salary_str).replace(",", "").replace("₹", "").replace(" ", "").strip()
        return float(cleaned)
    except (ValueError, TypeError):
        return 0.0


def _fmt(amount: float) -> str:
    return f"₹ {amount:,.2f}"


# ---------------------------------------------------------------------------
# Core: run→paragraph-level placeholder replacement
# ---------------------------------------------------------------------------

_PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")


def _replace_in_paragraph(para, context: dict) -> None:
    """
    Replace all {{key}} tokens inside a paragraph, even when Word has split
    the placeholder text across multiple Runs (which it frequently does).
    """
    full_text = "".join(run.text for run in para.runs)
    if "{{" not in full_text:
        return

    def _replacer(m: re.Match) -> str:
        return context.get(m.group(0), m.group(0))

    new_text = _PLACEHOLDER_RE.sub(_replacer, full_text)
    if new_text == full_text:
        return

    if not para.runs:
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _replace_in_doc(doc: Document, context: dict) -> None:
    """Apply _replace_in_paragraph to every paragraph in the document,
    including paragraphs inside table cells."""
    for para in doc.paragraphs:
        _replace_in_paragraph(para, context)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, context)


# ---------------------------------------------------------------------------
# Salary table builder
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _write_cell(
    cell,
    text: str,
    bold: bool = False,
    font_size: int = 9,
    color: Optional[str] = None,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _build_salary_table(doc: Document, annual_salary_str: str) -> object:
    """
    Build a 3-column salary breakdown table (Component | Per Month | Per Annum)
    using calculate_salary() and return its raw XML element for repositioning.
    No Remarks column; no Other Benefits section.
    """
    ctc = _parse_salary(annual_salary_str)
    if ctc <= 0:
        return None

    sal = calculate_salary("", ctc)   # employee_name not needed for table content

    R = WD_ALIGN_PARAGRAPH.RIGHT
    C = WD_ALIGN_PARAGRAPH.CENTER
    L = WD_ALIGN_PARAGRAPH.LEFT

    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    col_widths = [Inches(3.4), Inches(1.3), Inches(1.3)]

    def add_row(texts, bg=_C_WHITE, bolds=None, fg=None, aligns=None, fs=9):
        row    = table.add_row()
        bolds  = bolds  or [False] * 3
        fg     = fg     or [None]  * 3
        aligns = aligns or [L]     * 3
        for idx, (cell, text) in enumerate(zip(row.cells, texts)):
            _set_cell_bg(cell, bg)
            _write_cell(cell, text, bold=bolds[idx], font_size=fs,
                        color=fg[idx], align=aligns[idx])
            cell.width = col_widths[idx]

    W3 = [_C_WHITE] * 3

    # ── Header ───────────────────────────────────────────────────────────────
    add_row(
        ["Components of the Salary Package", "Amount (P.M.)", "Amount (P.A.)"],
        bg=_C_DARK_BLUE, bolds=[True]*3, fg=[_C_WHITE]*3,
        aligns=[L, R, R],
    )

    # ── Earnings section ─────────────────────────────────────────────────────
    add_row(["Basic Salary Components", "", ""],
            bg=_C_MID_BLUE, bolds=[True, False, False], fg=[_C_WHITE]*3)

    earnings = [
        ("Basic Salary",                       sal.basic_pm,           sal.basic_pa),
        ("House Rent Allowance (HRA)",          sal.hra_pm,             sal.hra_pa),
        ("Medical Allowances",                  sal.medical_pm,         sal.medical_pa),
        ("Transport Allowances",                sal.transport_pm,       sal.transport_pa),
        ("Deployment / Performance Allowances", sal.deployment_pm,      sal.deployment_pa),
        ("Fixed Allowances",                    sal.fixed_allowance_pm, sal.fixed_allowance_pa),
    ]
    for i, (name, pm, pa) in enumerate(earnings):
        bg = _C_LIGHT_BLUE if i % 2 == 0 else _C_WHITE
        add_row([name, f"{pm:,.2f}", f"{pa:,.2f}"], bg=bg, aligns=[L, R, R])

    # Gross total
    add_row(
        ["Gross / CTC  (A)", f"{sal.gross_pm:,.2f}", f"{sal.gross_pa:,.2f}"],
        bg=_C_MID_BLUE, bolds=[True, True, True],
        fg=[_C_WHITE]*3, aligns=[L, R, R],
    )

    # ── Deductions section ────────────────────────────────────────────────────
    add_row(["", "", ""])
    add_row(["Less Deductions:", "", ""],
            bg=_C_DARK_BLUE, bolds=[True, False, False], fg=[_C_WHITE]*3)

    deductions = [
        ("EPF — Employee Contribution (12%)", sal.epf_employee_pm, sal.epf_employee_pa),
        ("EPF — Employer Contribution (12%)", sal.epf_employer_pm, sal.epf_employer_pa),
        ("ESIC — Employee (0.75%)",           sal.esic_employee_pm, sal.esic_employee_pa),
    ]
    for i, (name, pm, pa) in enumerate(deductions):
        bg = _C_LIGHT_BLUE if i % 2 == 0 else _C_WHITE
        add_row([name, f"{pm:,.2f}", f"{pa:,.2f}"], bg=bg, aligns=[L, R, R])

    add_row(
        ["Total Deductions  (B)", f"{sal.total_deductions_pm:,.2f}", f"{sal.total_deductions_pa:,.2f}"],
        bg=_C_MID_BLUE, bolds=[True, True, True],
        fg=[_C_WHITE]*3, aligns=[L, R, R],
    )

    # ── Net income ────────────────────────────────────────────────────────────
    add_row(["", "", ""])
    add_row(
        ["Net Income  (C = A − B)", f"{sal.net_pm:,.2f}", f"{sal.net_pa:,.2f}"],
        bg=_C_GOLD, bolds=[True, True, True],
        fg=[_C_WHITE]*3, aligns=[L, R, R], fs=10,
    )

    return table._tbl   # raw XML element


def _inject_signature_image(
    doc: Document,
    hiring_manager_name: str,
    signature_img_bytes: Optional[bytes],
    width_inches: float = 1.5,
) -> None:
    """
    Find the paragraph containing {{Signature SignatureHRManager2}},
    clear the placeholder text, and insert the hiring manager's signature
    image inline.  Falls back to the plain name if the image is unavailable.
    """
    placeholder = "{{Signature SignatureHRManager2}}"

    def _process_para(para) -> bool:
        """Returns True if the placeholder was found and handled."""
        full_text = "".join(run.text for run in para.runs)
        if placeholder not in full_text:
            return False

        # Clear all runs so the raw placeholder text disappears
        for run in para.runs:
            run.text = ""

        if signature_img_bytes:
            # Add the image into the first run of this paragraph
            run = para.add_run()
            run.add_picture(io.BytesIO(signature_img_bytes), width=Inches(width_inches))
            logger.info(
                "offer_letter_generator — hiring manager signature image injected"
            )
        else:
            # Graceful fallback: write the name as plain text
            para.add_run(hiring_manager_name or "")
            logger.warning(
                "offer_letter_generator — signature image unavailable; "
                "falling back to plain-text name"
            )
        return True

    # Search top-level paragraphs
    for para in doc.paragraphs:
        if _process_para(para):
            return

    # Search inside table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _process_para(para):
                        return

    logger.warning(
        "offer_letter_generator — {{Signature SignatureHRManager2}} placeholder not found"
    )


def _inject_candidate_signature(
    doc: Document,
    candidate_name: str,
    signature_img_bytes: Optional[bytes],
    width_inches: float = 1.5,
) -> None:
    """
    Find **all** paragraphs containing ``{{Signature CandidateSignature}}``,
    clear the placeholder text, and insert the candidate's signature image
    inline in each one.  Falls back to a "Signed by <name>" label if no
    image is supplied.

    NOTE: The template may contain the placeholder more than once (e.g. once
    in the body and once in a signature table).  All occurrences are replaced.
    """
    placeholder = "{{Signature CandidateSignature}}"
    found_count = 0

    def _process_para(para) -> None:
        nonlocal found_count
        full_text = "".join(run.text for run in para.runs)
        if placeholder not in full_text:
            return

        # Clear all runs so the raw placeholder text disappears
        for run in para.runs:
            run.text = ""

        if signature_img_bytes:
            run = para.add_run()
            run.add_picture(io.BytesIO(signature_img_bytes), width=Inches(width_inches))
        else:
            # Fallback: candidate name as plain text
            para.add_run(f"Signed by: {candidate_name}" if candidate_name else "Candidate Signature")

        found_count += 1

    # Search top-level paragraphs — replace every match (do NOT return early)
    for para in doc.paragraphs:
        _process_para(para)

    # Search inside table cells — replace every match (do NOT return early)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_para(para)

    if found_count == 0:
        logger.warning(
            "offer_letter_generator — {{Signature CandidateSignature}} placeholder not found"
        )
    elif signature_img_bytes:
        logger.info(
            f"offer_letter_generator — candidate signature image injected "
            f"into {found_count} location(s)"
        )
    else:
        logger.warning(
            f"offer_letter_generator — candidate signature image unavailable; "
            f"fell back to plain-text label in {found_count} location(s)"
        )



def _inject_salary_table(doc: Document, annual_salary_str: str) -> None:
    """
    Find the paragraph containing _SALARY_TABLE_MARKER, build the salary table
    from the CTC, insert it at that position, and remove the marker paragraph.
    """
    # --- find the marker paragraph ----------------------------------------
    marker_para = None
    for para in doc.paragraphs:
        if _SALARY_TABLE_MARKER in para.text:
            marker_para = para
            break

    if marker_para is None:
        logger.warning("offer_letter_generator — salary table marker not found in document")
        return

    ctc = _parse_salary(annual_salary_str)
    if ctc <= 0:
        for run in marker_para.runs:
            run.text = ""
        logger.warning(f"offer_letter_generator — cannot parse salary '{annual_salary_str}'; table skipped")
        return

    # --- build table (appended at end of doc body then repositioned) ------
    tbl_xml = _build_salary_table(doc, annual_salary_str)
    if tbl_xml is None:
        return

    # --- reposition the table XML right after the marker paragraph --------
    para_xml = marker_para._element
    para_xml.addnext(tbl_xml)

    # --- remove the now-redundant placeholder paragraph -------------------
    para_xml.getparent().remove(para_xml)

    logger.info(
        f"offer_letter_generator — salary table injected (CTC={ctc:,.0f})"
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_context(
    *,
    first_name: str,
    last_name: str,
    job_title: str,
    department: str,
    location: str,
    offer_expire_date: str,
    joining_date,
    annual_salary: str,
    hiring_manager_name: str,
    letter_date=None,
) -> dict:
    """
    Build the {{placeholder}} → value mapping used to fill the template.

    NOTE: {{BasicInfo.FullSalaryStructureWithCTCDeductionsAndNetPay}} is mapped
    to an internal marker that is later replaced by a real salary table by
    _inject_salary_table().
    """
    if letter_date is None:
        letter_date = date.today()

    candidate_full_name = f"{first_name} {last_name}".strip()

    return {
        # Basic info
        "{{BasicInfo.FirstName}}":                                     first_name or "",
        "{{BasicInfo.LastName}}":                                      last_name or "",
        "{{BasicInfo.JobTitle}}":                                      job_title or "",
        "{{BasicInfo.FullName}}":                                      candidate_full_name or "",
        "{{BasicInfo.Department}}":                                    department or "",
        "{{BasicInfo.Location}}":                                      location or "",
        "{{BasicInfo.DateOfferExpires}}":                              offer_expire_date or "",
        "{{BasicInfo.MediumDate}}":                                    _fmt_date_medium(letter_date),
        "{{BasicInfo.DateJoiningMedium}}":                             _fmt_date_medium(joining_date),
        "{{BasicInfo.AnnualSalary}}":                                  annual_salary or "",
        "{{BasicInfo.FullSalaryWithBonusOthersWithoutDeduction}}":     annual_salary or "",
        # This placeholder → sentinel; replaced with a real table post-processing
        "{{BasicInfo.FullSalaryStructureWithCTCDeductionsAndNetPay}}": _SALARY_TABLE_MARKER,
        # Job info
        "{{JobInfo.Location}}":                                        location or "",
        # Signatures
        # NOTE: {{Signature SignatureHRManager2}} is intentionally omitted here;
        #       it is replaced with the actual PNG image by _inject_signature_image().
    }


def generate_filled_docx(
    *,
    first_name: str,
    last_name: str,
    job_title: str,
    department: str,
    location: str,
    offer_expire_date,
    joining_date,
    annual_salary: str,
    hiring_manager_name: str,
    template_path: Optional[str] = None,
    hm_signature_bytes: Optional[bytes] = None,
    candidate_name: Optional[str] = None,
    candidate_signature_bytes: Optional[bytes] = None,
) -> bytes:
    """
    Download the SharePoint template, fill all placeholders, inject the salary
    table, and return the resulting .docx as raw bytes.

    Args:
        template_path:            Override the default SHAREPOINT_TEMPLATE_PATH env var.
        hm_signature_bytes:       PNG bytes of the hiring manager's signature.
                                  When provided, the static SharePoint file is skipped.
                                  When None, the legacy SharePoint file is downloaded.
        candidate_name:           Candidate full name (used as fallback label if no sig image).
        candidate_signature_bytes: PNG bytes of the candidate's signature.
                                  When provided, the {{Signature CandidateSignature}}
                                  placeholder is replaced with the actual image.

    Returns:
        Bytes of the filled .docx file.
    """
    tpl_path = template_path or TEMPLATE_PATH
    logger.info(f"offer_letter_generator — downloading template: {tpl_path}")

    # 1. Fetch template
    raw_bytes = download_file(tpl_path)

    # 2. Open in python-docx
    doc = Document(io.BytesIO(raw_bytes))

    # 3. Build replacement context
    context = build_context(
        first_name=first_name,
        last_name=last_name,
        job_title=job_title,
        department=department,
        location=location,
        offer_expire_date=_fmt_date_medium(offer_expire_date),
        joining_date=joining_date,
        annual_salary=annual_salary,
        hiring_manager_name=hiring_manager_name,
    )

    logger.info(
        f"offer_letter_generator — replacing placeholders for "
        f"{first_name} {last_name} (position: {job_title})"
    )

    # 4. Resolve hiring-manager signature bytes
    #    Priority: caller-supplied bytes → SharePoint static file → plain-text fallback
    sig_img_bytes: Optional[bytes] = hm_signature_bytes
    if sig_img_bytes is None:
        try:
            sig_img_bytes = download_file(HR_MANAGER_SIGNATURE_PATH)
            logger.info(
                f"offer_letter_generator — static signature image downloaded "
                f"({len(sig_img_bytes)} bytes) from {HR_MANAGER_SIGNATURE_PATH}"
            )
        except Exception as exc:
           logger.error(f"Error: {str(exc)}", exc_info=True)
            logger.warning(
                f"offer_letter_generator — could not download static signature image "
                f"from '{HR_MANAGER_SIGNATURE_PATH}': {exc}"
            )

    # 5. Text placeholder replacement (sentinel written for salary table)
    _replace_in_doc(doc, context)

    # 6. Inject hiring-manager signature image
    _inject_signature_image(doc, hiring_manager_name, sig_img_bytes)

    # 7. Inject candidate signature image (if supplied, otherwise placeholder already handled in context)
    if candidate_signature_bytes is not None or candidate_name:
        _inject_candidate_signature(doc, candidate_name or "", candidate_signature_bytes)

    # 8. Inject the salary breakdown table where the sentinel now sits
    _inject_salary_table(doc, annual_salary)

    # 9. Serialise to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    logger.info("offer_letter_generator — document generated successfully")
    return output.read()


def generated_file_path(candidate_id: str, offer_id: int) -> str:
    """Return the SharePoint path where the generated offer letter will be saved."""
    return f"{GENERATED_FOLDER}/{candidate_id}/offer_{offer_id}.docx"


def signed_offer_file_path(candidate_id: str, offer_id: int) -> str:
    """Return the SharePoint path where the fully-signed offer letter will be saved."""
    return f"{GENERATED_FOLDER}/{candidate_id}/offer_{offer_id}_signed.docx"


def inject_candidate_signature_into_docx(
    docx_bytes: bytes,
    candidate_name: str,
    signature_img_bytes: Optional[bytes],
    width_inches: float = 1.5,
) -> bytes:
    """
    Take an already-generated `.docx` (e.g. the HM-approved version stored on
    SharePoint) and inject the candidate's signature image into the
    ``{{Signature CandidateSignature}}`` placeholder **without re-generating
    the whole document from the template**.

    This preserves the hiring-manager signature that was already embedded
    during the approval step.

    Args:
        docx_bytes:          Raw bytes of the existing `.docx` file.
        candidate_name:      Candidate full name (used as fallback label if no
                             signature image is available).
        signature_img_bytes: PNG bytes of the candidate's hand-written signature.
        width_inches:        Width of the signature image in the document.

    Returns:
        Bytes of the updated `.docx` with the candidate signature injected.
    """
    doc = Document(io.BytesIO(docx_bytes))
    _inject_candidate_signature(doc, candidate_name, signature_img_bytes, width_inches)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    logger.info("offer_letter_generator — candidate signature injected into existing docx")
    return output.read()
